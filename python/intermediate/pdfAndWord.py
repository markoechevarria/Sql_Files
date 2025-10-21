import PyPDF2
pdfFileObj = open('meetingminutes.pdf', 'rb')
pdfReader = PyPDF2.PdfReader(pdfFileObj)
print(len(pdfReader.pages))

pageObj = pdfReader.pages[0]
print(pageObj.extract_text())

pdfFileObj.close()



### Decrypting PDFs

"""
pdfReader = PyPDF2.PdfReader(open('encrypted.pdf', 'rb'))
print(pdfReader.is_encrypted)
print(pdfReader.pages[0])
"""
pdfReader = PyPDF2.PdfReader(open('encrypted.pdf', 'rb'))
pdfReader.decrypt('rosebud')
pdfObj = pdfReader.pages[0]
print(pdfObj.extract_text())

### Creating PDFs

"""
	PdfWriter
"""

### Copying PDFs

pdf1File = open('meetingminutes.pdf', 'rb')
pdf2File = open('meetingminutes2.pdf', 'rb')

pdf1Reader = PyPDF2.PdfReader(pdf1File)
pdf2Reader = PyPDF2.PdfReader(pdf2File)

pdfWriter = PyPDF2.PdfWriter()

for pageNum in range(len(pdf1Reader.pages)):
	pageObj = pdf1Reader.pages[pageNum]
	pdfWriter.add_page(pageObj)

for pageNum in range(len(pdf2Reader.pages)):
	pageObj = pdf2Reader.pages[pageNum]
	pdfWriter.add_page(pageObj)

pdfOutputFile = open('combinedminutes.pdf', 'wb')
pdfWriter.write(pdfOutputFile)
pdfOutputFile.close()
pdf1File.close()
pdf2File.close()


### Rotating Pages

"""
   >>> import PyPDF2
   >>> minutesFile = open('meetingminutes.pdf', 'rb')
   >>> pdfReader = PyPDF2.PdfFileReader(minutesFile)
➊ >>> page = pdfReader.getPage(0)
➋ >>> page.rotateClockwise(90)
   {'/Contents': [IndirectObject(961, 0), IndirectObject(962, 0),
   --snip--
   }
   >>> pdfWriter = PyPDF2.PdfFileWriter()
   >>> pdfWriter.addPage(page)
➌ >>> resultPdfFile = open('rotatedPage.pdf', 'wb')
   >>> pdfWriter.write(resultPdfFile)
   >>> resultPdfFile.close()
   >>> minutesFile.close()


"""

### Overlaying pages

"""
import PyPDF2
minutesFile = open('meetingminutes.pdf', 'rb')
pdfReader = PyPDF2.PdfFileReader(minutesFile)
minutesFirstPage = pdfReader.getPage(0)
pdfWatermarkReader = PyPDF2.PdfFileReader(open('watermark.pdf', 'rb'))
minutesFirstPage.mergePage(pdfWatermarkReader.getPage(0))
pdfWriter = PyPDF2.PdfFileWriter()
pdfWriter.addPage(minutesFirstPage)

for pageNum in range(1, pdfReader.numPages):
	pageObj = pdfReader.getPage(pageNum)
	pdfWriter.addPage(pageObj)

resultPdfFile = open('watermarkedCover.pdf', 'wb')
pdfWriter.write(resultPdfFile)
minutesFile.close()
resultPdfFile.close()
"""

### Encrypting PDFs

"""
import PyPDF2
pdfFile = open('meetingminutes.pdf', 'rb')
pdfReader = PyPDF2.PdfFileReader(pdfFile)
pdfWriter = PyPDF2.PdfFileWriter()
for pageNum in range(pdfReader.numPages):
	pdfWriter.addPage(pdfReader.getPage(pageNum))

pdfWriter.encrypt('swordfish')
resultPdf = open('encryptedminutes.pdf', 'wb')
pdfWriter.write(resultPdf)
resultPdf.close()
"""

###  Find all pdf Files
"""
import PyPDF2, os
pdfFiles = []
for filename in os.listdir('.'):
	if filename.endswith('.pdf'):
		pdfFiles.append(filename)
pdfFiles.sort(key = str.lower)
pdfWriter = PyPDF2.PdfWriter()

for filename in pdfFiles:
	pdfFileObj = open(filename, 'rb')
	pdfReader = PyPDF2.PdfReader(pdfFileObj)

for pageNum in range(1, len(pdfReader.pages)):
	pageObj = pdfReader.page(pageNUm)
	pdfWriter.add_page(pageObj)

pdfOutput = open('allminutes.pdf', 'wb')
pdfWriter.write(pdfOutput)
pdfOutput.close()
"""

##3 Reading word documents

import docx
doc = docx.Document('demo.docx')
print(len(doc.paragraphs))

print(doc.paragraphs[0].text)
print(len(doc.paragraphs[1].runs))

print(doc.paragraphs[1].runs[0].text)

### Getting the fulltext from a .docx file

def getText(filename):
	doc = docx.Document(filename)
	fullText = []
	for para in doc.paragraphs:
		fullText.append(para.text)
	return '\n'.join(fullText)
print(getText('demo.docx'))

"""
	doc.paragraphs[0].style = 'Normal'
	doc.paragraphs[1].runs[0].underline = True
	doc.paragraphs[1].runs[0].style = 'QuoteChar'
"""

### Writing word documents
import docx
doc = docx.Document()
doc.add_paragraph('Hello world')
doc.save('helloworld.docx')

doc.add_paragraph('Hello world2')
paraObj1 = doc.add_paragraph('This is a second paragraph')
paraObj2 = doc.add_paragraph('This is a yet another paragraph')
paraObj1.add_run('This is being added to the second paragraph')
doc.save('multipleParagraphs.docx')

doc.add_heading('Header 0', 0)
doc.add_heading('Header 1', 1)
doc.add_heading('Header 2', 2)
doc.add_heading('Header 3', 3)
doc.add_heading('Header 4', 4)

doc.add_paragraph('This is on the first page')
doc.paragraphs[0].runs[0].add_break(docx.enum.text.WD_BREAK.PAGE)
doc.add_paragraph('this is on the second page')
""" doc.save('twopage.docx') """


doc.add_picture('zophie.png', width=docx.shared.Inches(1), height = docx.shared.Cm(4))
